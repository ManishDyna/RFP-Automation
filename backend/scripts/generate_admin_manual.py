"""
Bahra Electric - RFP Portal Admin User Manual Generator
Generates a professional .docx admin manual with branded formatting.
Reuses helper functions from generate_user_manual.py.

Usage:
    python scripts/generate_admin_manual.py
"""

import os
import sys

# Add scripts directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from generate_user_manual import (
    # Constants
    PROJECT_ROOT, LOGO_PATH,
    BAHRA_DARK_GRAY, BAHRA_RED, BAHRA_LIGHT_GRAY, WHITE,
    FONT_HEADING, FONT_BODY,
    FONT_SIZE_BODY, FONT_SIZE_H1, FONT_SIZE_H2, FONT_SIZE_H3,
    # Helpers
    configure_styles, setup_headers_footers,
    add_heading, add_para, add_bullet, add_numbered, add_tip, add_important,
    add_styled_table, insert_toc_field, add_page_number_field,
    set_cell_shading, set_table_borders,
)
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Admin-specific Metadata ───────────────────────────────────────────────────
ADMIN_DOC_TITLE = "Admin User Manual"
ADMIN_DOC_SUBTITLE = "Bahra Electric Industrial Company"
ADMIN_DOC_VERSION = "1.0"
ADMIN_DOC_DATE = "February 2026"
ADMIN_DOC_NUMBER = "BE-RFP-AM-001"
ADMIN_OUTPUT_DOCX = os.path.join(PROJECT_ROOT, "docs", "Bahra_Electric_RFP_Portal_Admin_Manual.docx")


# =============================================================================
# FRONT MATTER
# =============================================================================

def build_cover_page(doc):
    """Build the admin manual cover page."""
    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Logo
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO_PATH):
        run = p_logo.add_run()
        run.add_picture(LOGO_PATH, width=Inches(3))

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Red divider
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), 'CF2E2E')
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(ADMIN_DOC_TITLE)
    run_title.font.name = FONT_HEADING
    run_title.font.size = Pt(36)
    run_title.font.color.rgb = BAHRA_DARK_GRAY
    run_title.bold = True
    p_title.paragraph_format.space_after = Pt(8)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(ADMIN_DOC_SUBTITLE)
    run_sub.font.name = FONT_BODY
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = BAHRA_LIGHT_GRAY
    p_sub.paragraph_format.space_after = Pt(12)

    # Version & Date
    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ver = p_ver.add_run(f"Version {ADMIN_DOC_VERSION}  |  {ADMIN_DOC_DATE}")
    run_ver.font.name = FONT_BODY
    run_ver.font.size = Pt(12)
    run_ver.font.color.rgb = BAHRA_LIGHT_GRAY
    p_ver.paragraph_format.space_after = Pt(6)

    # Audience
    p_aud = doc.add_paragraph()
    p_aud.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_aud = p_aud.add_run("Audience: System Administrators")
    run_aud.font.name = FONT_BODY
    run_aud.font.size = Pt(11)
    run_aud.font.color.rgb = BAHRA_LIGHT_GRAY

    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # Confidential badge
    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_conf = p_conf.add_run("CONFIDENTIAL")
    run_conf.font.name = FONT_BODY
    run_conf.font.size = Pt(14)
    run_conf.font.color.rgb = BAHRA_RED
    run_conf.bold = True

    doc.add_page_break()


def build_confidentiality_notice(doc):
    """Build the confidentiality notice page."""
    add_heading(doc, "Confidentiality Notice", level=1)

    add_para(doc, (
        "This document is the property of Bahra Electric Industrial Company and contains "
        "proprietary and confidential information. It is intended solely for use by authorized "
        "system administrators of Bahra Electric. Unauthorized reproduction, distribution, or "
        "disclosure of this document, in whole or in part, is strictly prohibited."
    ))

    add_para(doc, (
        "The information contained herein includes sensitive system administration procedures, "
        "permission configurations, and security-related features. Recipients of this document "
        "are responsible for maintaining its confidentiality and must not share it with any "
        "third party without prior written consent from Bahra Electric Industrial Company."
    ))

    add_para(doc, (
        "If you have received this document in error, please notify the IT Department at Bahra Electric "
        "immediately and destroy all copies in your possession."
    ))

    doc.add_page_break()


def build_document_control(doc):
    """Build the document control page."""
    add_heading(doc, "Document Control", level=1)

    add_heading(doc, "Document Information", level=2)
    add_styled_table(doc,
        headers=["Field", "Details"],
        rows=[
            ["Document Title", ADMIN_DOC_TITLE],
            ["Document Number", ADMIN_DOC_NUMBER],
            ["Version", ADMIN_DOC_VERSION],
            ["Classification", "Confidential"],
            ["Effective Date", ADMIN_DOC_DATE],
            ["Owner", "Bahra Electric \u2013 IT Department"],
        ],
        col_widths=[Inches(2.5), Inches(4)]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    add_heading(doc, "Revision History", level=2)
    add_styled_table(doc,
        headers=["Version", "Date", "Author", "Description"],
        rows=[
            ["1.0", ADMIN_DOC_DATE, "IT Department", "Initial release"],
        ],
        col_widths=[Inches(1), Inches(1.5), Inches(1.5), Inches(2.5)]
    )

    doc.add_page_break()


def build_table_of_contents(doc):
    """Build the Table of Contents page."""
    add_heading(doc, "Table of Contents", level=1)
    insert_toc_field(doc)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    p_note = doc.add_paragraph()
    run_note = p_note.add_run(
        "Note: To update this table of contents in Microsoft Word, right-click anywhere in "
        "the table above and select 'Update Field', then choose 'Update entire table'."
    )
    run_note.font.name = FONT_BODY
    run_note.font.size = Pt(9)
    run_note.font.color.rgb = BAHRA_LIGHT_GRAY
    run_note.italic = True

    doc.add_page_break()


# =============================================================================
# CONTENT SECTIONS
# =============================================================================

def build_section_1(doc):
    """Section 1: Introduction"""
    add_heading(doc, "1. Introduction", level=1)

    add_para(doc, (
        "This Admin User Manual provides comprehensive guidance for system administrators "
        "of the Bahra Electric RFP Portal. It covers all administrative features including "
        "user management, role and permission configuration, audit trail monitoring, SAP "
        "credential management, and automation scheduling."
    ))

    add_heading(doc, "1.1 Admin vs End-User Access", level=2)
    add_para(doc, (
        "The RFP Portal uses a Role-Based Access Control (RBAC) system. Features visible to "
        "each user depend on their assigned role and permissions. Administrators have access to "
        "all portal features, while end-users (RFP Bidders) see only the core RFP management features."
    ))

    add_styled_table(doc,
        headers=["Feature", "Admin", "RFP Bidder"],
        rows=[
            ["Dashboard", "\u2713", "\u2713"],
            ["RFP Insights", "\u2713", "\u2713"],
            ["Material Insights", "\u2713", "\u2713"],
            ["Activity Logs", "\u2713", "\u2713"],
            ["Submit / Decline / Download RFPs", "\u2713", "\u2713"],
            ["Analytics", "\u2713", "\u2717"],
            ["User Management", "\u2713", "\u2717"],
            ["Role Management", "\u2713", "\u2717"],
            ["Audit Logs", "\u2713", "\u2717"],
            ["SAP Password Logs", "\u2713", "\u2717"],
            ["Schedule Automation", "\u2713", "\u2717"],
            ["SAP Password Change", "\u2713", "\u2717"],
        ]
    )

    add_heading(doc, "1.2 Prerequisites", level=2)
    add_para(doc, "Before using the administrative features, ensure that:")
    add_bullet(doc, "You have an active account with the Admin role assigned.")
    add_bullet(doc, "You can access the RFP Portal URL in Google Chrome (latest version recommended).")
    add_bullet(doc, "You have reviewed the End-User Manual (BE-RFP-UM-001) for familiarity with core portal features.")

    add_heading(doc, "1.3 Document Conventions", level=2)
    add_bullet(doc, " text indicates a clickable UI element (button, menu item, or link).", bold_prefix="Bold")
    add_bullet(doc, " text indicates user-entered values or field names.", bold_prefix="Italic")
    add_bullet(doc, ' prefixed paragraphs provide helpful shortcuts or best practices.', bold_prefix='"Tip:" ')
    add_bullet(doc, ' prefixed paragraphs highlight critical warnings or restrictions.', bold_prefix='"Important:" ')

    doc.add_page_break()


def build_section_2(doc):
    """Section 2: Admin Navigation"""
    add_heading(doc, "2. Admin Navigation", level=1)

    add_para(doc, (
        "When logged in with an Admin role, the left sidebar displays additional menu items "
        "under the Administration section. These items are only visible if your role has the "
        "corresponding permissions."
    ))

    add_heading(doc, "2.1 Administration Menu", level=2)
    add_para(doc, (
        "The Administration section appears below the main menu items in the sidebar and contains:"
    ))

    add_styled_table(doc,
        headers=["Menu Item", "Permission Required", "Description"],
        rows=[
            ["Users", "user_management.view", "Manage user accounts \u2013 create, edit, delete, activate/deactivate, and unlock"],
            ["Roles", "role_management.view", "Manage roles and assign granular permissions"],
            ["Audit Logs", "audit_logs.view", "View the complete audit trail of system events"],
            ["Analytics", "analytics.view", "View analytics dashboard with charts and drill-downs"],
            ["SAP Logs", "sap_password.view", "View history of SAP password changes"],
        ]
    )

    add_heading(doc, "2.2 Settings Section", level=2)
    add_para(doc, (
        "Below the Quick Actions in the sidebar, the Settings section provides access to:"
    ))

    add_styled_table(doc,
        headers=["Setting", "Permission Required", "Description"],
        rows=[
            ["Schedule", "schedule_automation.manage", "Configure automated RFP download schedules"],
            ["SAP Password", "sap_password.change", "Update SAP portal credentials used for automation"],
        ]
    )

    add_important(doc, (
        "If you do not see these menu items, your role may not have the required permissions. "
        "Contact another administrator to update your role."
    ))

    doc.add_page_break()


def build_section_3(doc):
    """Section 3: User Management"""
    add_heading(doc, "3. User Management", level=1)
    add_para(doc, (
        "The User Management page allows you to create, edit, and manage all user accounts in the system. "
        "Navigate to Users under the Administration section in the sidebar."
    ))

    # 3.1 User List
    add_heading(doc, "3.1 User List", level=2)
    add_para(doc, "The main table displays all registered users with the following columns:")

    add_styled_table(doc,
        headers=["Column", "Description"],
        rows=[
            ["Name", "User\u2019s full name"],
            ["Email", "User\u2019s email address (used for login)"],
            ["Mobile", "Contact phone number (displays \u201c\u2013\u201d if not provided)"],
            ["Role", "Assigned role (e.g., Admin, RFP Bidder). Admin role shows a blue badge."],
            ["Status", "Account status: Active (green), Inactive (gray), or Locked (red with lock icon)"],
            ["Created", "Account creation date"],
            ["Actions", "Action buttons (Edit, Activate/Deactivate, Unlock, Delete)"],
        ]
    )

    add_para(doc, (
        "Use the Search users... box at the top to filter the list by name or email. "
        "The search updates results in real-time as you type."
    ))

    # 3.2 Create a New User
    add_heading(doc, "3.2 Create a New User", level=2)
    add_para(doc, "To create a new user account:", bold=False)
    add_numbered(doc, 'Click the Add User button (top-right corner). Requires user_management.create permission.')
    add_numbered(doc, 'The "Add New User" dialog opens with the following fields:')

    add_styled_table(doc,
        headers=["Field", "Required", "Validation", "Notes"],
        rows=[
            ["Name", "Yes", "Minimum 2 characters", "Full name of the user"],
            ["Email", "Yes", "Must be a valid email format", "Used as the login username"],
            ["Mobile", "No", "None", "Optional contact number"],
            ["Role", "Yes", "Must select a role", "Dropdown showing all active roles"],
            ["Password", "Yes", "Minimum 6 characters", "Initial login password"],
        ]
    )

    add_numbered(doc, "Fill in all required fields and click Create (or Save) to create the account.")
    add_numbered(doc, 'A success notification "User created successfully" confirms the action.')

    # 3.3 Edit a User
    add_heading(doc, "3.3 Edit a User", level=2)
    add_numbered(doc, "Click the pencil icon in the Actions column of the user you want to edit.")
    add_numbered(doc, 'The "Edit User" dialog opens with the user\u2019s current information pre-filled.')
    add_numbered(doc, "Modify the desired fields. Note:")
    add_bullet(doc, " is read-only and cannot be changed after creation.", bold_prefix="Email")
    add_bullet(doc, " can be left blank to keep the current password unchanged.", bold_prefix="Password")
    add_numbered(doc, "Click Save to apply changes.")

    # 3.4 User Status Management
    add_heading(doc, "3.4 User Status Management", level=2)
    add_para(doc, "User accounts can be in one of three states:")

    add_styled_table(doc,
        headers=["Status", "Indicator", "Description"],
        rows=[
            ["Active", "Green badge", "User can log in and use the portal normally"],
            ["Inactive", "Gray badge", "User account is deactivated and cannot log in"],
            ["Locked", "Red badge with lock icon", "Account locked due to 5 failed login attempts within 5 minutes"],
        ]
    )

    add_heading(doc, "Deactivate a User", level=3)
    add_para(doc, (
        "Click the shield-off icon (orange) in the Actions column. This immediately prevents "
        "the user from logging in. The user\u2019s data is preserved and the account can be "
        "reactivated at any time. Requires user_management.activate permission."
    ))

    add_heading(doc, "Activate a User", level=3)
    add_para(doc, (
        "Click the shield-check icon (green) in the Actions column for an inactive user. "
        "This restores login access. Requires user_management.activate permission."
    ))

    add_heading(doc, "Unlock a User", level=3)
    add_para(doc, (
        "When a user enters the wrong password 5 times within 5 minutes, their account is "
        "automatically locked for 30 minutes. To unlock immediately, click the unlock icon "
        "(blue) in the Actions column. Requires user_management.edit permission."
    ))

    # 3.5 Delete a User
    add_heading(doc, "3.5 Delete a User", level=2)
    add_numbered(doc, "Click the trash icon (red) in the Actions column. Requires user_management.delete permission.")
    add_numbered(doc, 'A confirmation dialog appears: "Are you sure you want to delete {name}? This action cannot be undone."')
    add_numbered(doc, "Click Delete to confirm, or Cancel to abort.")

    add_important(doc, "Deleting a user is permanent and cannot be undone. Consider deactivating the user instead if you may need to restore access later.")

    doc.add_page_break()


def build_section_4(doc):
    """Section 4: Role Management"""
    add_heading(doc, "4. Role Management", level=1)
    add_para(doc, (
        "The Role Management page allows you to create custom roles and assign granular permissions. "
        "Navigate to Roles under the Administration section in the sidebar."
    ))

    # 4.1 Role List
    add_heading(doc, "4.1 Role List", level=2)
    add_para(doc, "The table displays all roles with the following columns:")

    add_styled_table(doc,
        headers=["Column", "Description"],
        rows=[
            ["Role Name", "Name of the role"],
            ["Description", "Brief description of the role\u2019s purpose"],
            ["Permissions", "Badge showing the number of assigned permissions"],
            ["Type", "System (blue badge) or Custom (outline badge)"],
            ["Status", "Active (green) or Inactive (red)"],
            ["Actions", "Edit and Delete buttons (conditional on permissions)"],
        ]
    )

    add_para(doc, (
        "Use the Search roles... box to filter roles by name or description."
    ))

    add_heading(doc, "System vs Custom Roles", level=3)
    add_styled_table(doc,
        headers=["Attribute", "System Roles", "Custom Roles"],
        rows=[
            ["Examples", "Admin, RFP Bidder", "Any role you create"],
            ["Can be edited", "Yes (permissions can be changed)", "Yes"],
            ["Can be deleted", "No (delete button is hidden)", "Yes"],
            ["Badge", "Blue \u201cSystem\u201d badge", "Outline \u201cCustom\u201d badge"],
        ]
    )

    # 4.2 Create a New Role
    add_heading(doc, "4.2 Create a New Role", level=2)
    add_numbered(doc, "Click the Create Role button (top-right). Requires role_management.create permission.")
    add_numbered(doc, 'The "Create Role" dialog opens with three sections:')

    add_para(doc, "Role Details:", bold=True, space_after=Pt(4))
    add_styled_table(doc,
        headers=["Field", "Required", "Validation", "Notes"],
        rows=[
            ["Name", "Yes", "Minimum 2 characters", 'e.g., "Manager", "Viewer"'],
            ["Description", "No", "None", "Brief description of the role\u2019s purpose"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_para(doc, "Permissions:", bold=True, space_after=Pt(4))
    add_para(doc, (
        "Permissions are organized into collapsible groups by module. Each group displays a "
        "counter (e.g., 3/5) showing how many permissions are selected out of the total available."
    ))
    add_bullet(doc, "Click a group header to expand or collapse it.")
    add_bullet(doc, 'Use the "Select All" / "Deselect All" button within each group to toggle all permissions at once.')
    add_bullet(doc, "Check or uncheck individual permissions using the checkboxes.")
    add_bullet(doc, 'The header shows the total count: "Permissions (X selected)".')

    add_numbered(doc, "Select the desired permissions for this role.")
    add_numbered(doc, 'Click Create Role to save. A success notification "Role created successfully" confirms the action.')

    # 4.3 Edit a Role
    add_heading(doc, "4.3 Edit a Role", level=2)
    add_numbered(doc, "Click the pencil icon in the Actions column. Requires role_management.edit permission.")
    add_numbered(doc, 'The "Edit Role" dialog opens with the current name, description, and permissions pre-loaded.')
    add_numbered(doc, "Modify the role name, description, or permissions as needed.")
    add_numbered(doc, "Click Save to apply changes.")

    add_tip(doc, "You can edit system roles (Admin, RFP Bidder) to adjust their permissions, but you cannot delete them.")

    # 4.4 Delete a Role
    add_heading(doc, "4.4 Delete a Role", level=2)
    add_numbered(doc, "Click the trash icon in the Actions column. Requires role_management.delete permission.")
    add_numbered(doc, 'A confirmation dialog appears: "Are you sure you want to delete the role \\"{name}\\"? This will deactivate the role. Users with this role may lose access."')
    add_numbered(doc, "Click Delete to confirm.")

    add_important(doc, (
        "Deleting a role deactivates it. Users currently assigned to this role may lose access "
        "to permission-gated features. Reassign affected users to another role before deleting."
    ))

    doc.add_page_break()


def build_section_5(doc):
    """Section 5: Audit Logs"""
    add_heading(doc, "5. Audit Logs", level=1)
    add_para(doc, (
        "The Audit Logs page provides a complete audit trail of all security and administrative "
        "events in the system. Navigate to Audit Logs under the Administration section in the sidebar. "
        "Requires audit_logs.view permission."
    ))

    # 5.1 Filter Options
    add_heading(doc, "5.1 Filter Options", level=2)
    add_para(doc, "Use the filter bar at the top to narrow down audit log entries:")

    add_styled_table(doc,
        headers=["Filter", "Type", "Description"],
        rows=[
            ["Category", "Dropdown", "Filter by event category: AUTH, USER, ROLE, RFP, SYSTEM"],
            ["Action", "Dropdown", "Filter by specific action (e.g., LOGIN, USER_CREATED, ROLE_UPDATED)"],
            ["Actor Email", "Text input", "Filter by the email of the user who performed the action"],
            ["From", "Date picker", "Show only events after this date"],
            ["To", "Date picker", "Show only events before this date"],
        ]
    )

    add_para(doc, (
        "Click Clear Filters to reset all filters. Filters can be combined \u2014 "
        "for example, filter by Category = USER and a date range to see all user management "
        "actions within a specific period."
    ))

    # 5.2 Audit Log Table
    add_heading(doc, "5.2 Audit Log Table", level=2)
    add_para(doc, "Each audit log entry displays:")

    add_styled_table(doc,
        headers=["Column", "Description"],
        rows=[
            ["Timestamp", "Date and time when the event occurred"],
            ["Category", "Color-coded badge indicating the event category"],
            ["Action", "The specific action performed (e.g., LOGIN, USER_CREATED)"],
            ["Actor", "Name and email of the user who performed the action"],
            ["Target", "The type and ID of the affected entity (e.g., User: john@example.com)"],
            ["Details", "Additional context or changes made (JSON formatted)"],
            ["IP Address", "The source IP address of the action (monospace font)"],
        ]
    )

    # 5.3 Category Colors
    add_heading(doc, "5.3 Category Color Coding", level=2)

    add_styled_table(doc,
        headers=["Category", "Color", "Events Tracked"],
        rows=[
            ["AUTH", "Blue", "Login, logout, failed login attempts, password changes and resets"],
            ["USER", "Green", "User creation, updates, deletion, activation, deactivation, unlocking"],
            ["ROLE", "Purple", "Role creation, updates, deletion, permission changes"],
            ["RFP", "Orange", "RFP-related administrative actions"],
            ["SYSTEM", "Gray", "System-level events (e.g., role seeding, configuration changes)"],
        ]
    )

    # 5.4 Pagination
    add_heading(doc, "5.4 Pagination", level=2)
    add_para(doc, (
        "Audit logs are paginated. The footer shows the current page, total pages, and total "
        "record count. Use the Previous and Next buttons to navigate between pages. "
        "Changing any filter resets the view to page 1."
    ))

    doc.add_page_break()


def build_section_6(doc):
    """Section 6: SAP Password Management"""
    add_heading(doc, "6. SAP Password Management", level=1)
    add_para(doc, (
        "The SAP Password feature allows administrators to update the SAP portal credentials "
        "used by the automation system. These credentials are required for submitting and "
        "declining RFPs on external portals."
    ))

    add_heading(doc, "6.1 Change SAP Password", level=2)
    add_numbered(doc, "In the sidebar, find the Settings section and click SAP Password. Requires sap_password.change permission.")
    add_numbered(doc, 'The "Change SAP Password" dialog opens with three fields:')

    add_styled_table(doc,
        headers=["Field", "Required", "Validation", "Notes"],
        rows=[
            ["Username", "Yes", "Cannot be empty", "The SAP portal username"],
            ["Enter Password", "Yes", "Cannot be empty", "The new SAP portal password"],
            ["Confirm Password", "Yes", "Must match the password above", "For your own confirmation only"],
        ]
    )

    add_numbered(doc, "Fill in all fields and click Save.")
    add_numbered(doc, 'A success notification "SAP password updated successfully" confirms the change.')

    add_important(doc, (
        "The SAP password is encrypted before being stored. Changing this password affects "
        "all automation operations (Submit, Decline) that interact with SAP-integrated portals."
    ))

    doc.add_page_break()


def build_section_7(doc):
    """Section 7: SAP Password Logs"""
    add_heading(doc, "7. SAP Password Logs", level=1)
    add_para(doc, (
        "The SAP Password Logs page displays a history of all SAP password changes. "
        "Navigate to SAP Logs under the Administration section. Requires sap_password.view permission."
    ))

    add_heading(doc, "7.1 Log Table", level=2)
    add_para(doc, "The table shows the following information for each password change:")

    add_styled_table(doc,
        headers=["Column", "Description"],
        rows=[
            ["Username", "The SAP username whose password was changed"],
            ["Password", "The password value (masked by default with \u2022\u2022\u2022\u2022\u2022\u2022\u2022\u2022)"],
            ["Changed By", "The email of the administrator who made the change"],
            ["Created", "Date and time of the password change"],
            ["Status", 'Always shows "Saved" with a green badge'],
        ]
    )

    add_heading(doc, "7.2 Password Visibility Controls", level=2)
    add_para(doc, "Each password cell has two controls:")
    add_bullet(doc, " \u2013 Click the eye icon to reveal the password. Click the eye-off icon to hide it again.", bold_prefix="Show/Hide Toggle")
    add_bullet(doc, " \u2013 Click the copy icon to copy the password to your clipboard. A notification confirms the copy.", bold_prefix="Copy to Clipboard")

    add_para(doc, 'Use the "Search logs..." box to filter the log entries by any field (username, email, date).')

    doc.add_page_break()


def build_section_8(doc):
    """Section 8: Schedule Automation"""
    add_heading(doc, "8. Schedule Automation", level=1)
    add_para(doc, (
        "The Schedule Automation feature allows you to configure recurring automated RFP downloads. "
        "In the sidebar, find the Settings section and click Schedule. "
        "Requires schedule_automation.manage permission."
    ))

    add_heading(doc, "8.1 Schedule Configuration", level=2)
    add_para(doc, 'The "Schedule Automation" dialog provides the following settings:')

    add_styled_table(doc,
        headers=["Field", "Required", "Default", "Description"],
        rows=[
            ["Interval", "Yes", "6", "How often the automation should run (numeric value)"],
            ["Frequency", "Yes", "Hour", "Unit of time: Minute, Hour, Day, Week, or Month"],
            ["Time Zone", "Yes", "Asia/Kolkata", "The timezone for scheduling calculations"],
            ["Start Time", "No", "\u2013", "When the first run should start (date and time picker)"],
        ]
    )

    add_heading(doc, "8.2 Available Time Zones", level=2)
    add_styled_table(doc,
        headers=["Time Zone", "Display Label"],
        rows=[
            ["Asia/Kolkata", "(UTC+05:30) Chennai, Kolkata, Mumbai, New Delhi"],
            ["Asia/Riyadh", "(UTC+03:00) Kuwait, Riyadh"],
            ["UTC", "(UTC+00:00) Coordinated Universal Time"],
            ["Europe/London", "(UTC+00:00) Dublin, Edinburgh, Lisbon, London"],
            ["Europe/Berlin", "(UTC+01:00) Amsterdam, Berlin, Rome, Stockholm"],
        ]
    )

    add_heading(doc, "8.3 Advanced Options", level=2)
    add_para(doc, (
        'Click "Show advanced options" to reveal additional settings:'
    ))

    add_styled_table(doc,
        headers=["Field", "Default", "Description"],
        rows=[
            ["Max Concurrency", "1", "Maximum number of automation tasks that can run simultaneously"],
            ["Notes", "\u2013", "Optional note describing the purpose of this schedule"],
        ]
    )

    add_heading(doc, "8.4 Saving the Schedule", level=2)
    add_numbered(doc, "Configure the desired interval, frequency, and timezone.")
    add_numbered(doc, "Optionally set a start time and advanced options.")
    add_numbered(doc, "Click Save Schedule.")
    add_numbered(doc, 'A success notification "Schedule saved successfully" confirms the configuration.')

    add_tip(doc, (
        "The default schedule runs every 6 hours. For high-volume environments, consider running "
        "every 1\u20132 hours. For less active environments, daily or weekly may be sufficient."
    ))

    doc.add_page_break()


def build_section_9(doc):
    """Section 9: Permission Reference"""
    add_heading(doc, "9. Permission Reference", level=1)
    add_para(doc, (
        "The RFP Portal uses a granular permission system with 22 individual permissions "
        "organized into 10 modules. Each permission controls access to a specific feature or action."
    ))

    add_heading(doc, "9.1 Complete Permission List", level=2)

    # User Management
    add_heading(doc, "User Management", level=3)
    add_styled_table(doc,
        headers=["Permission Key", "Label", "Description"],
        rows=[
            ["user_management.view", "View user list", "Access the User Management page and see all users"],
            ["user_management.create", "Create new users", "Create new user accounts via the Add User dialog"],
            ["user_management.edit", "Edit existing users", "Modify user details and unlock locked accounts"],
            ["user_management.delete", "Delete users", "Permanently remove user accounts"],
            ["user_management.activate", "Activate/deactivate users", "Toggle user account active/inactive status"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Role Management
    add_heading(doc, "Role Management", level=3)
    add_styled_table(doc,
        headers=["Permission Key", "Label", "Description"],
        rows=[
            ["role_management.view", "View roles and permissions", "Access the Role Management page"],
            ["role_management.create", "Create new roles", "Create custom roles with permissions"],
            ["role_management.edit", "Edit roles and assign permissions", "Modify role details and permissions"],
            ["role_management.delete", "Delete roles", "Deactivate custom roles"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # SAP Password
    add_heading(doc, "SAP Password", level=3)
    add_styled_table(doc,
        headers=["Permission Key", "Label", "Description"],
        rows=[
            ["sap_password.view", "View SAP password logs", "Access the SAP Password Logs page"],
            ["sap_password.change", "Change SAP password", "Update SAP portal credentials"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Schedule & Automation
    add_heading(doc, "Schedule and Automation", level=3)
    add_styled_table(doc,
        headers=["Permission Key", "Label", "Description"],
        rows=[
            ["schedule_automation.view", "View automation schedules", "View the current schedule configuration"],
            ["schedule_automation.manage", "Create/edit/delete schedules", "Configure and modify automation schedules"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Analytics, Dashboard, Logs, Audit, Material, RFP
    add_heading(doc, "Portal Features", level=3)
    add_styled_table(doc,
        headers=["Permission Key", "Label", "Description"],
        rows=[
            ["analytics.view", "View analytics dashboard", "Access the Analytics page with charts"],
            ["dashboard.view", "View main dashboard", "Access the main Dashboard page"],
            ["rfp.view", "View RFP insights and details", "Access the RFP Insights page"],
            ["rfp.download", "Download RFPs from portal", "Trigger RFP downloads from company portals"],
            ["rfp.submit", "Submit RFPs", "Submit RFPs via the Submit RFP dialog"],
            ["rfp.decline", "Decline RFPs", "Decline RFPs via the Decline RFP dialog"],
            ["logs.view", "View automation activity logs", "Access the Activity Logs page"],
            ["audit_logs.view", "View audit trail logs", "Access the Audit Logs page"],
            ["material_insights.view", "View material insights", "Access the Material Insights page"],
        ]
    )

    doc.add_page_break()


def build_section_10(doc):
    """Section 10: Troubleshooting & FAQ"""
    add_heading(doc, "10. Troubleshooting and FAQ", level=1)

    add_heading(doc, "I don\u2019t see the Administration menu items", level=2)
    add_bullet(doc, (
        "Administration menu items (Users, Roles, Audit Logs, SAP Logs) are only visible if "
        "your role has the corresponding view permissions."
    ))
    add_bullet(doc, (
        "Ask another administrator to check your role and ensure it includes the required permissions "
        "(e.g., user_management.view, role_management.view, audit_logs.view)."
    ))

    add_heading(doc, "I cannot create or edit users", level=2)
    add_bullet(doc, (
        "Creating users requires the user_management.create permission. Editing requires user_management.edit."
    ))
    add_bullet(doc, "Verify your role has these specific permissions in the Role Management page.")

    add_heading(doc, "A user is locked out and cannot log in", level=2)
    add_bullet(doc, (
        "Accounts are automatically locked after 5 failed login attempts within 5 minutes. "
        "The lockout lasts 30 minutes."
    ))
    add_bullet(doc, (
        "To unlock immediately: Go to Users, find the locked user (red lock icon), and click "
        "the blue unlock icon in the Actions column."
    ))

    add_heading(doc, "I accidentally deleted a user", level=2)
    add_bullet(doc, (
        "User deletion is permanent. The user\u2019s account and data cannot be recovered."
    ))
    add_bullet(doc, (
        "To restore access, create a new account for the user with the same email and role. "
        "Historical activity logs referencing the old account will still be preserved."
    ))
    add_bullet(doc, (
        "To avoid accidental deletion, consider deactivating users instead of deleting them."
    ))

    add_heading(doc, "I cannot delete a system role (Admin or RFP Bidder)", level=2)
    add_bullet(doc, (
        "System roles cannot be deleted. They are built-in roles required for portal operation."
    ))
    add_bullet(doc, (
        "You can edit system roles to add or remove permissions, but the delete button is "
        "intentionally hidden for these roles."
    ))

    add_heading(doc, "Users lost access after I deleted a custom role", level=2)
    add_bullet(doc, (
        "Deleting a role deactivates it. Users assigned to that role may lose access to "
        "permission-gated features."
    ))
    add_bullet(doc, (
        "Reassign affected users to an active role via the User Management page (edit each user "
        "and select a new role)."
    ))

    add_heading(doc, "The SAP password change is not taking effect", level=2)
    add_bullet(doc, (
        "After changing the SAP password, the new credentials are used for all subsequent "
        "automation runs (Submit, Decline)."
    ))
    add_bullet(doc, (
        "If automation still fails, verify the credentials are correct by checking the SAP Password "
        "Logs page and confirming the username and password match the portal."
    ))

    add_heading(doc, "The scheduled automation is not running", level=2)
    add_bullet(doc, (
        "Verify the schedule is configured correctly by clicking Schedule in the sidebar Settings."
    ))
    add_bullet(doc, (
        "Check that the start time (if set) has passed and the interval/frequency are reasonable."
    ))
    add_bullet(doc, (
        "Review Activity Logs to see if automation runs are being triggered and whether they "
        "are completing or failing."
    ))

    add_heading(doc, "How do I see who made a specific change?", level=2)
    add_bullet(doc, (
        "Navigate to Audit Logs. Use the filters to search by action type, category, actor email, "
        "or date range."
    ))
    add_bullet(doc, (
        "The Actor column shows who performed each action, and the Details column shows what was changed."
    ))

    add_heading(doc, "How do I create a read-only admin role?", level=2)
    add_bullet(doc, (
        "Create a new role and assign only the view permissions: user_management.view, "
        "role_management.view, audit_logs.view, sap_password.view, schedule_automation.view."
    ))
    add_bullet(doc, (
        "This allows the user to see all admin pages without the ability to create, edit, or delete anything."
    ))

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_para(doc, "For additional support, contact the IT Department.", bold=True)

    doc.add_page_break()


def build_appendix_a(doc):
    """Appendix A: Default Role Permissions Matrix"""
    add_heading(doc, "Appendix A: Default Role Permissions", level=1)
    add_para(doc, (
        "The following table shows the permissions assigned to each default system role. "
        "These can be customized via the Role Management page."
    ))

    add_styled_table(doc,
        headers=["Permission", "Admin", "RFP Bidder"],
        rows=[
            ["user_management.view", "\u2713", "\u2717"],
            ["user_management.create", "\u2713", "\u2717"],
            ["user_management.edit", "\u2713", "\u2717"],
            ["user_management.delete", "\u2713", "\u2717"],
            ["user_management.activate", "\u2713", "\u2717"],
            ["role_management.view", "\u2713", "\u2717"],
            ["role_management.create", "\u2713", "\u2717"],
            ["role_management.edit", "\u2713", "\u2717"],
            ["role_management.delete", "\u2713", "\u2717"],
            ["sap_password.view", "\u2713", "\u2717"],
            ["sap_password.change", "\u2713", "\u2717"],
            ["schedule_automation.view", "\u2713", "\u2717"],
            ["schedule_automation.manage", "\u2713", "\u2717"],
            ["analytics.view", "\u2713", "\u2717"],
            ["audit_logs.view", "\u2713", "\u2717"],
            ["dashboard.view", "\u2713", "\u2713"],
            ["rfp.view", "\u2713", "\u2713"],
            ["rfp.download", "\u2713", "\u2713"],
            ["rfp.submit", "\u2713", "\u2713"],
            ["rfp.decline", "\u2713", "\u2713"],
            ["logs.view", "\u2713", "\u2713"],
            ["material_insights.view", "\u2713", "\u2713"],
        ],
        col_widths=[Inches(3.5), Inches(1.5), Inches(1.5)]
    )

    doc.add_page_break()


def build_appendix_b(doc):
    """Appendix B: Audit Log Actions Reference"""
    add_heading(doc, "Appendix B: Audit Log Actions Reference", level=1)
    add_para(doc, (
        "The following table lists all actions tracked in the audit log, organized by category."
    ))

    add_styled_table(doc,
        headers=["Category", "Action", "Description"],
        rows=[
            ["AUTH", "LOGIN", "User successfully logged into the portal"],
            ["AUTH", "LOGOUT", "User logged out of the portal"],
            ["AUTH", "LOGIN_FAILED", "Failed login attempt (wrong password or unknown email)"],
            ["AUTH", "PASSWORD_CHANGED", "User changed their own password via Profile Settings"],
            ["AUTH", "PASSWORD_RESET", "User reset their password via the Forgot Password flow"],
            ["USER", "USER_CREATED", "Administrator created a new user account"],
            ["USER", "USER_UPDATED", "Administrator edited a user\u2019s name, mobile, role, or password"],
            ["USER", "USER_DELETED", "Administrator permanently deleted a user account"],
            ["USER", "USER_ACTIVATED", "Administrator activated a previously deactivated user"],
            ["USER", "USER_DEACTIVATED", "Administrator deactivated an active user account"],
            ["USER", "USER_UNLOCKED", "Administrator unlocked a locked user account"],
            ["ROLE", "ROLE_CREATED", "Administrator created a new role"],
            ["ROLE", "ROLE_UPDATED", "Administrator updated a role\u2019s name or description"],
            ["ROLE", "ROLE_DELETED", "Administrator deleted (deactivated) a custom role"],
            ["ROLE", "ROLE_PERMISSIONS_UPDATED", "Administrator changed the permissions assigned to a role"],
            ["SYSTEM", "SEED_ROLES", "System seeded default roles (Admin, RFP Bidder) during initial setup"],
        ],
        col_widths=[Inches(1.2), Inches(2.3), Inches(3)]
    )

    doc.add_page_break()


def build_back_cover(doc):
    """Build the back cover page."""
    for _ in range(8):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO_PATH):
        run = p_logo.add_run()
        run.add_picture(LOGO_PATH, width=Inches(2.5))

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(ADMIN_DOC_SUBTITLE)
    run_name.font.name = FONT_HEADING
    run_name.font.size = Pt(16)
    run_name.font.color.rgb = BAHRA_DARK_GRAY
    run_name.bold = True

    p_copy = doc.add_paragraph()
    p_copy.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_copy = p_copy.add_run("\u00A9 2026 Bahra Electric. All rights reserved.")
    run_copy.font.name = FONT_BODY
    run_copy.font.size = Pt(10)
    run_copy.font.color.rgb = BAHRA_LIGHT_GRAY


# =============================================================================
# MAIN ORCHESTRATOR
# =============================================================================

def generate_admin_manual():
    """Main entry point to generate the admin user manual."""
    print("=" * 60)
    print("  Bahra Electric - Admin User Manual Generator")
    print("=" * 60)

    # Create document
    print("\n[1/3] Building document structure...")
    doc = Document()

    # Configure page layout
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.75)

    # Configure styles
    configure_styles(doc)

    # Build all sections
    print("  Building cover page...")
    build_cover_page(doc)

    print("  Building front matter...")
    build_confidentiality_notice(doc)
    build_document_control(doc)
    build_table_of_contents(doc)

    print("  Building Section 1: Introduction...")
    build_section_1(doc)

    print("  Building Section 2: Admin Navigation...")
    build_section_2(doc)

    print("  Building Section 3: User Management...")
    build_section_3(doc)

    print("  Building Section 4: Role Management...")
    build_section_4(doc)

    print("  Building Section 5: Audit Logs...")
    build_section_5(doc)

    print("  Building Section 6: SAP Password Management...")
    build_section_6(doc)

    print("  Building Section 7: SAP Password Logs...")
    build_section_7(doc)

    print("  Building Section 8: Schedule Automation...")
    build_section_8(doc)

    print("  Building Section 9: Permission Reference...")
    build_section_9(doc)

    print("  Building Section 10: Troubleshooting & FAQ...")
    build_section_10(doc)

    print("  Building Appendix A: Default Role Permissions...")
    build_appendix_a(doc)

    print("  Building Appendix B: Audit Log Actions...")
    build_appendix_b(doc)

    print("  Building Back Cover...")
    build_back_cover(doc)

    # Setup headers and footers
    print("\n[2/3] Configuring headers and footers...")
    # Override the title for admin manual in header/footer
    for i, section in enumerate(doc.sections):
        if i == 0:
            section.different_first_page_header_footer = True
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.clear()
            run = hp.add_run(ADMIN_DOC_TITLE)
            run.font.name = FONT_BODY
            run.font.size = Pt(9)
            run.font.color.rgb = BAHRA_LIGHT_GRAY
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.clear()
            run_left = fp.add_run("Bahra Electric  |  CONFIDENTIAL")
            run_left.font.name = FONT_BODY
            run_left.font.size = Pt(8)
            run_left.font.color.rgb = BAHRA_LIGHT_GRAY

            tab_stops = fp.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.5), alignment=2)
            fp.add_run("\t")
            run_page_label = fp.add_run("Page ")
            run_page_label.font.name = FONT_BODY
            run_page_label.font.size = Pt(8)
            run_page_label.font.color.rgb = BAHRA_LIGHT_GRAY
            add_page_number_field(fp)

    # Save
    print("\n[3/3] Saving document...")
    os.makedirs(os.path.dirname(ADMIN_OUTPUT_DOCX), exist_ok=True)
    doc.save(ADMIN_OUTPUT_DOCX)
    print(f"  Saved to: {ADMIN_OUTPUT_DOCX}")

    print("\n" + "=" * 60)
    print("  DONE! Open the document in Microsoft Word.")
    print("  Right-click the Table of Contents > 'Update Field'")
    print("  to populate page numbers.")
    print("=" * 60)


if __name__ == "__main__":
    generate_admin_manual()

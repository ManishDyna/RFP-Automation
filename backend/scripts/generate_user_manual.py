"""
Bahra Electric - RFP Portal User Manual Generator
Generates a professional .docx user manual from scratch with branded formatting.
Extracts images from the original user-manual.docx and rebuilds with proper structure.

Usage:
    python scripts/generate_user_manual.py
"""

import os
import sys
import zipfile
import shutil
import tempfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────────
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ORIGINAL_DOCX = os.path.join(os.path.expanduser("~"), "Downloads", "user-manual.docx")
OUTPUT_DOCX = os.path.join(PROJECT_ROOT, "docs", "Bahra_Electric_RFP_Portal_User_Manual.docx")
LOGO_PATH = os.path.join(PROJECT_ROOT, "static", "images", "bahra-electric-logo.png")

# ── Brand Colors ───────────────────────────────────────────────────────────────
BAHRA_DARK_GRAY = RGBColor(0x32, 0x37, 0x3C)
BAHRA_RED = RGBColor(0xCF, 0x2E, 0x2E)
BAHRA_RED_DARK = RGBColor(0xA8, 0x24, 0x24)
BAHRA_LIGHT_GRAY = RGBColor(0xAB, 0xB8, 0xC3)
TABLE_HEADER_BG = "A82424"
TABLE_ALT_ROW_BG = "F2F2F2"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── Typography ─────────────────────────────────────────────────────────────────
FONT_HEADING = "Cambria"
FONT_BODY = "Calibri"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_H1 = Pt(24)
FONT_SIZE_H2 = Pt(16)
FONT_SIZE_H3 = Pt(13)
FONT_SIZE_TABLE = Pt(10)
FONT_SIZE_CAPTION = Pt(9)

# ── Document Metadata ─────────────────────────────────────────────────────────
DOC_TITLE = "RFP Portal User Manual"
DOC_SUBTITLE = "Bahra Electric Industrial Company"
DOC_VERSION = "1.0"
DOC_DATE = "February 2026"
DOC_NUMBER = "BE-RFP-UM-001"

# ── Figure counter ─────────────────────────────────────────────────────────────
_figure_counter = 0


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def extract_images_from_docx(docx_path):
    """Extract all images from word/media/ inside the .docx ZIP."""
    images = {}
    with zipfile.ZipFile(docx_path, 'r') as z:
        for name in z.namelist():
            if name.startswith('word/media/') and not name.endswith('/'):
                fname = os.path.basename(name)
                images[fname] = z.read(name)
    return images


def set_cell_shading(cell, color_hex):
    """Apply background color to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table, color="ABB8C3", size="4"):
    """Apply borders to an entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Remove existing borders
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:color'), color)
        el.set(qn('w:space'), '0')
        borders.append(el)
    tblPr.append(borders)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a branded table with red header row and alternating shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = FONT_SIZE_TABLE
        run.font.name = FONT_BODY
        run.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_cell_shading(cell, TABLE_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = FONT_SIZE_TABLE
            run.font.name = FONT_BODY
            run.font.color.rgb = BAHRA_DARK_GRAY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_ROW_BG)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

    set_table_borders(table)
    return table


def add_figure(doc, image_path, caption_text, width=Inches(6.5), centered=True):
    """Insert an image with a numbered caption below it."""
    global _figure_counter
    _figure_counter += 1

    if not os.path.exists(image_path):
        p = doc.add_paragraph(f"[Image not found: {os.path.basename(image_path)}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    # Image paragraph
    p_img = doc.add_paragraph()
    if centered:
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(image_path, width=width)

    # Caption paragraph
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run(f"Figure {_figure_counter}: {caption_text}")
    run_cap.italic = True
    run_cap.font.size = FONT_SIZE_CAPTION
    run_cap.font.color.rgb = BAHRA_LIGHT_GRAY
    run_cap.font.name = FONT_BODY


def add_heading(doc, text, level=1):
    """Add a heading with proper styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_HEADING
        if level == 1:
            run.font.size = FONT_SIZE_H1
            run.font.color.rgb = BAHRA_DARK_GRAY
        elif level == 2:
            run.font.size = FONT_SIZE_H2
            run.font.color.rgb = BAHRA_DARK_GRAY
        elif level == 3:
            run.font.size = FONT_SIZE_H3
            run.font.color.rgb = BAHRA_RED
    return h


def add_para(doc, text, bold=False, italic=False, space_after=Pt(6)):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = FONT_SIZE_BODY
    run.font.color.rgb = BAHRA_DARK_GRAY
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point. Optionally bold the prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.name = FONT_BODY
        run_b.font.size = FONT_SIZE_BODY
        run_b.font.color.rgb = BAHRA_DARK_GRAY
        run_rest = p.add_run(text)
        run_rest.font.name = FONT_BODY
        run_rest.font.size = FONT_SIZE_BODY
        run_rest.font.color.rgb = BAHRA_DARK_GRAY
    else:
        p.clear()
        run = p.add_run(text)
        run.font.name = FONT_BODY
        run.font.size = FONT_SIZE_BODY
        run.font.color.rgb = BAHRA_DARK_GRAY
    p.paragraph_format.space_after = Pt(3)
    return p


def add_numbered(doc, text, bold_prefix=None):
    """Add a numbered list item."""
    p = doc.add_paragraph(style='List Number')
    p.clear()
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.name = FONT_BODY
        run_b.font.size = FONT_SIZE_BODY
        run_b.font.color.rgb = BAHRA_DARK_GRAY
        run_rest = p.add_run(text)
        run_rest.font.name = FONT_BODY
        run_rest.font.size = FONT_SIZE_BODY
        run_rest.font.color.rgb = BAHRA_DARK_GRAY
    else:
        run = p.add_run(text)
        run.font.name = FONT_BODY
        run.font.size = FONT_SIZE_BODY
        run.font.color.rgb = BAHRA_DARK_GRAY
    p.paragraph_format.space_after = Pt(3)
    return p


def add_tip(doc, text):
    """Add a tip/note paragraph with bold prefix."""
    p = doc.add_paragraph()
    run_prefix = p.add_run("Tip: ")
    run_prefix.bold = True
    run_prefix.italic = True
    run_prefix.font.name = FONT_BODY
    run_prefix.font.size = FONT_SIZE_BODY
    run_prefix.font.color.rgb = BAHRA_RED
    run_text = p.add_run(text)
    run_text.italic = True
    run_text.font.name = FONT_BODY
    run_text.font.size = FONT_SIZE_BODY
    run_text.font.color.rgb = BAHRA_DARK_GRAY
    p.paragraph_format.space_after = Pt(6)
    return p


def add_important(doc, text):
    """Add an important/warning paragraph."""
    p = doc.add_paragraph()
    run_prefix = p.add_run("Important: ")
    run_prefix.bold = True
    run_prefix.font.name = FONT_BODY
    run_prefix.font.size = FONT_SIZE_BODY
    run_prefix.font.color.rgb = BAHRA_RED
    run_text = p.add_run(text)
    run_text.font.name = FONT_BODY
    run_text.font.size = FONT_SIZE_BODY
    run_text.font.color.rgb = BAHRA_DARK_GRAY
    p.paragraph_format.space_after = Pt(6)
    return p


def insert_toc_field(doc):
    """Insert a TOC field code that Word will populate on Update Fields."""
    paragraph = doc.add_paragraph()
    # Begin field
    run1 = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fldChar1)
    # Field instruction
    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)
    # Separate
    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)
    # Placeholder text
    run4 = paragraph.add_run("Right-click here and select 'Update Field' to generate Table of Contents")
    run4.font.color.rgb = BAHRA_LIGHT_GRAY
    run4.font.italic = True
    run4.font.name = FONT_BODY
    run4.font.size = FONT_SIZE_BODY
    # End field
    run5 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)
    return paragraph


def add_page_number_field(paragraph):
    """Insert a PAGE field into a paragraph for page numbering."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run("1")
    run4.font.name = FONT_BODY
    run4.font.size = Pt(9)

    run5 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)


def configure_styles(doc):
    """Configure document-wide default styles."""
    styles = doc.styles

    # Normal
    normal = styles['Normal']
    normal.font.name = FONT_BODY
    normal.font.size = FONT_SIZE_BODY
    normal.font.color.rgb = BAHRA_DARK_GRAY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = FONT_HEADING
    h1.font.size = FONT_SIZE_H1
    h1.font.color.rgb = BAHRA_DARK_GRAY
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = FONT_HEADING
    h2.font.size = FONT_SIZE_H2
    h2.font.color.rgb = BAHRA_DARK_GRAY
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = FONT_HEADING
    h3.font.size = FONT_SIZE_H3
    h3.font.color.rgb = BAHRA_RED
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)


def setup_headers_footers(doc):
    """Configure running headers and footers for all sections after the cover."""
    for i, section in enumerate(doc.sections):
        if i == 0:
            # First section (cover + front matter): different first page, no header/footer on first page
            section.different_first_page_header_footer = True
            # First page header/footer are blank by default
            # Default header for remaining pages in section
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hp.clear()
            run = hp.add_run(DOC_TITLE)
            run.font.name = FONT_BODY
            run.font.size = Pt(9)
            run.font.color.rgb = BAHRA_LIGHT_GRAY
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.clear()
            run_left = fp.add_run("Bahra Electric  |  CONFIDENTIAL")
            run_left.font.name = FONT_BODY
            run_left.font.size = Pt(8)
            run_left.font.color.rgb = BAHRA_LIGHT_GRAY

            # Add tab stop for right-aligned page number
            tab_stops = fp.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.5), alignment=2)  # RIGHT alignment
            fp.add_run("\t")
            run_page_label = fp.add_run("Page ")
            run_page_label.font.name = FONT_BODY
            run_page_label.font.size = Pt(8)
            run_page_label.font.color.rgb = BAHRA_LIGHT_GRAY
            add_page_number_field(fp)
        else:
            section.header.is_linked_to_previous = True
            section.footer.is_linked_to_previous = True


# =============================================================================
# FRONT MATTER BUILDERS
# =============================================================================

def build_cover_page(doc, logo_path):
    """Build the cover page with logo, title, and branding."""
    # Spacer
    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    # Logo
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_path):
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Inches(3))

    # Spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Red divider line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), 'CF2E2E')
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(DOC_TITLE)
    run_title.font.name = FONT_HEADING
    run_title.font.size = Pt(36)
    run_title.font.color.rgb = BAHRA_DARK_GRAY
    run_title.bold = True
    p_title.paragraph_format.space_after = Pt(8)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(DOC_SUBTITLE)
    run_sub.font.name = FONT_BODY
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = BAHRA_LIGHT_GRAY
    p_sub.paragraph_format.space_after = Pt(12)

    # Version & Date
    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ver = p_ver.add_run(f"Version {DOC_VERSION}  |  {DOC_DATE}")
    run_ver.font.name = FONT_BODY
    run_ver.font.size = Pt(12)
    run_ver.font.color.rgb = BAHRA_LIGHT_GRAY
    p_ver.paragraph_format.space_after = Pt(6)

    # Audience
    p_aud = doc.add_paragraph()
    p_aud.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_aud = p_aud.add_run("Audience: End Users (RFP Bidders)")
    run_aud.font.name = FONT_BODY
    run_aud.font.size = Pt(11)
    run_aud.font.color.rgb = BAHRA_LIGHT_GRAY

    # Spacer
    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # Confidential badge
    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_conf = p_conf.add_run("CONFIDENTIAL")
    run_conf.font.name = FONT_BODY
    run_conf.font.size = Pt(14)
    run_conf.font.color.rgb = BAHRA_RED
    run_conf.bold = True

    doc.add_page_break()


def build_confidentiality_notice(doc):
    """Build the confidentiality notice page."""
    add_heading(doc, "Confidentiality Notice", level=1)

    add_para(doc, (
        "This document is the property of Bahra Electric Industrial Company and contains "
        "proprietary and confidential information. It is intended solely for use by authorized "
        "personnel of Bahra Electric and its designated partners. Unauthorized reproduction, "
        "distribution, or disclosure of this document, in whole or in part, is strictly prohibited."
    ))

    add_para(doc, (
        "The information contained herein is provided on a need-to-know basis. Recipients of this "
        "document are responsible for maintaining its confidentiality and must not share it with "
        "any third party without prior written consent from Bahra Electric Industrial Company."
    ))

    add_para(doc, (
        "If you have received this document in error, please notify the IT Department at Bahra Electric "
        "immediately and destroy all copies in your possession. Any unauthorized use of this document "
        "may result in legal action."
    ))

    doc.add_page_break()


def build_document_control(doc):
    """Build the document control / revision history page."""
    add_heading(doc, "Document Control", level=1)

    # Document Information table
    add_heading(doc, "Document Information", level=2)
    add_styled_table(doc,
        headers=["Field", "Details"],
        rows=[
            ["Document Title", DOC_TITLE],
            ["Document Number", DOC_NUMBER],
            ["Version", DOC_VERSION],
            ["Classification", "Confidential"],
            ["Effective Date", DOC_DATE],
            ["Owner", "Bahra Electric \u2013 IT Department"],
        ],
        col_widths=[Inches(2.5), Inches(4)]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Revision History table
    add_heading(doc, "Revision History", level=2)
    add_styled_table(doc,
        headers=["Version", "Date", "Author", "Description"],
        rows=[
            ["1.0", DOC_DATE, "IT Department", "Initial release"],
        ],
        col_widths=[Inches(1), Inches(1.5), Inches(1.5), Inches(2.5)]
    )

    doc.add_page_break()


def build_table_of_contents(doc):
    """Build the Table of Contents page with a TOC field."""
    add_heading(doc, "Table of Contents", level=1)
    insert_toc_field(doc)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    p_note = doc.add_paragraph()
    run_note = p_note.add_run(
        "Note: To update this table of contents in Microsoft Word, right-click anywhere in "
        "the table above and select 'Update Field', then choose 'Update entire table'."
    )
    run_note.font.name = FONT_BODY
    run_note.font.size = Pt(9)
    run_note.font.color.rgb = BAHRA_LIGHT_GRAY
    run_note.italic = True

    doc.add_page_break()


# =============================================================================
# CONTENT SECTION BUILDERS
# =============================================================================

def build_section_1(doc):
    """Section 1: Introduction"""
    add_heading(doc, "1. Introduction", level=1)

    add_para(doc, (
        "The Bahra Electric RFP Portal is a web-based system that automates the management of "
        "Requests for Proposal (RFPs) across multiple supplier portals. It handles downloading RFPs "
        "from company portals, matching materials and keywords, submitting or declining RFPs, and "
        "sending email notifications \u2014 all from a single dashboard."
    ))

    add_para(doc, "What you can do with this portal:", bold=True)
    add_bullet(doc, "View all downloaded RFPs organized by company and status")
    add_bullet(doc, "Submit or decline RFPs with automated portal interaction")
    add_bullet(doc, "Track material and keyword matches for each RFP")
    add_bullet(doc, "Monitor automation runs and review activity logs")
    add_bullet(doc, "View analytics and participation reports")
    add_bullet(doc, "Download RFP Excel files for offline review")
    add_bullet(doc, "Receive email notifications with interactive Adaptive Cards in Outlook")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_para(doc, "Supported Companies:", bold=True)
    add_bullet(doc, "Saudi Energy")
    add_bullet(doc, "Aramco e-Marketplace")
    add_bullet(doc, "SABIC \u2013 Saudi Basic Industries Corp.")
    add_bullet(doc, "HADEED \u2013 RAJHI STEEL")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_para(doc, "Browser Requirements:", bold=True)
    add_para(doc, "Google Chrome (latest version recommended) is the supported browser for the RFP Portal.")

    # Document Conventions subsection
    add_heading(doc, "1.1 Document Conventions", level=2)
    add_para(doc, "The following conventions are used throughout this manual:")
    add_bullet(doc, " text indicates a clickable user interface element (button, menu item, or link).", bold_prefix="Bold")
    add_bullet(doc, " text indicates user-entered values or field names.", bold_prefix="Italic")
    add_bullet(doc, ' prefixed paragraphs provide helpful shortcuts or best practices.', bold_prefix='"Tip:" ')
    add_bullet(doc, ' prefixed paragraphs highlight critical warnings or restrictions.', bold_prefix='"Important:" ')
    add_bullet(doc, "Figures are numbered sequentially and captioned below each screenshot.")

    doc.add_page_break()


def build_section_2(doc, image_paths):
    """Section 2: Getting Started"""
    add_heading(doc, "2. Getting Started", level=1)

    # 2.1 Logging In
    add_heading(doc, "2.1 Logging In", level=2)
    add_numbered(doc, "Open the RFP Portal URL in your browser.")
    add_numbered(doc, 'You will see the Login page with the Bahra Electric logo and "RFP Automation System" heading.')
    add_numbered(doc, "Enter your Email address in the email field.")
    add_numbered(doc, "Enter your Password in the password field.")
    add_numbered(doc, "Optionally, check Remember me to stay signed in.")
    add_numbered(doc, "Click Sign In.")

    add_para(doc, "After logging in, you will be redirected to the Dashboard page.")

    if "image1.png" in image_paths:
        add_figure(doc, image_paths["image1.png"], "Login Page")

    # 2.2 Forgot Password
    add_heading(doc, "2.2 Forgot Password", level=2)
    add_para(doc, "If you forget your password, follow these steps:")
    add_numbered(doc, 'On the Login page, click "Forgot password?" (located to the right of the Remember me checkbox).')
    add_numbered(doc, "A dialog will appear asking for your email address.")
    add_numbered(doc, "Enter the email associated with your account.")
    add_numbered(doc, "Click Send Reset Link.")
    add_numbered(doc, "Check your email inbox for a password reset link.")
    add_numbered(doc, "Click the link in the email and follow the instructions to set a new password.")

    if "image2.png" in image_paths:
        add_figure(doc, image_paths["image2.png"], "Forgot Password Dialog")

    # 2.3 Navigation Overview
    add_heading(doc, "2.3 Navigation Overview", level=2)
    add_para(doc, "After logging in, you will see two main areas:")

    add_para(doc, "Left Sidebar", bold=True, space_after=Pt(4))
    add_para(doc, "Your main navigation panel containing menu items, quick actions, and an automation status indicator.")

    add_styled_table(doc,
        headers=["Section", "Items", "Description"],
        rows=[
            ["Menu", "Dashboard", "Main overview with RFP metrics and management"],
            ["", "RFP Insights", "Detailed RFP data with advanced filters"],
            ["", "Material Insights", "Material and keyword matching analysis"],
            ["", "Activity Logs", "Automation run history and details"],
            ["Quick Actions", "Download RFPs", "Trigger RFP download from company portals"],
            ["", "Submit RFP", "Submit an RFP with file uploads"],
            ["", "Decline RFP", "Decline participation in an RFP"],
            ["Status Footer", "Automation Status", "Shows whether automation is Ready or Running"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_para(doc, (
        "The sidebar can be collapsed by clicking the arrow icon at the top-right corner of the sidebar, "
        "and expanded again by clicking the expand arrow. When collapsed, hovering over icons shows tooltip labels."
    ))

    add_para(doc, "Top Header", bold=True, space_after=Pt(4))
    add_para(doc, (
        "Displays the page title, description, and page-specific action buttons. "
        "The header also includes a user profile dropdown with options to view your profile or log out."
    ))

    if "image3.png" in image_paths:
        add_figure(doc, image_paths["image3.png"], "Portal Layout with Expanded Sidebar")
    if "image4.png" in image_paths:
        add_figure(doc, image_paths["image4.png"], "Collapsed Sidebar", width=Inches(0.6))

    doc.add_page_break()


def build_section_3(doc, image_paths):
    """Section 3: Dashboard"""
    add_heading(doc, "3. Dashboard", level=1)
    add_para(doc, "The Dashboard is your home page and provides a quick overview of all RFP activity.")

    # 3.1 Metric Cards
    add_heading(doc, "3.1 Metric Cards", level=2)
    add_para(doc, "At the top of the page, you will see four summary cards:")

    add_styled_table(doc,
        headers=["Card", "Description"],
        rows=[
            ["Total Downloaded RFPs", "Total number of RFPs downloaded from all portals"],
            ["Submitted", "Number of RFPs that have been submitted"],
            ["Declined", "Number of RFPs that have been declined"],
            ["Last Automation", "Timestamp of the last automation run"],
        ]
    )

    add_tip(doc, (
        "Clicking on the Total, Submitted, or Declined cards navigates to the RFP Insights page "
        "with the corresponding filter pre-applied."
    ))

    if "image5.png" in image_paths:
        add_figure(doc, image_paths["image5.png"], "Dashboard Metric Cards")

    # 3.2 RFP Management Section
    add_heading(doc, "3.2 RFP Management Section", level=2)
    add_para(doc, "Below the metrics, you will find the main RFP Management area, organized in two levels of tabs:")

    add_para(doc, (
        "Company Tabs (top row): Each tab represents a company (e.g., Saudi Energy, "
        "Aramco e-Marketplace). A badge next to each company name shows the total number of active RFPs."
    ))

    add_para(doc, "Status Sub-Tabs (second row): Within each company tab, RFPs are split by status:", bold=False)

    add_styled_table(doc,
        headers=["Status Tab", "Badge Color", "Meaning"],
        rows=[
            ["Open", "Amber", "RFPs awaiting action \u2013 can be submitted or declined"],
            ["Submitted", "Green", "RFPs already submitted to the portal"],
            ["Draft", "Gray", "RFPs saved as draft on the portal"],
            ["Declined", "Red", "RFPs where participation was declined"],
        ]
    )

    if "image6.png" in image_paths:
        add_figure(doc, image_paths["image6.png"], "RFP Management Section")

    # 3.3 RFP Table
    add_heading(doc, "3.3 RFP Table", level=2)
    add_para(doc, "Each status tab shows a table with the following columns:")

    add_styled_table(doc,
        headers=["Column", "Description"],
        rows=[
            ["RFP ID", "Unique identifier. Click to open the RFP on the external portal (opens in a new tab)."],
            ["Owner", "The RFP owner name from the portal."],
            ["Published", "Date the RFP was published."],
            ["Deadline", "RFP submission deadline date."],
            ["Match %", "Material match percentage \u2013 how well this RFP matches your company\u2019s materials."],
            ["Status", "Current status badge (Open, Submitted, Draft, Declined)."],
            ["Actions", "Action buttons available for this RFP."],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_heading(doc, "Match % Column Details", level=3)
    add_para(doc, "The Match % shows a color-coded badge and a mini progress bar:")

    add_styled_table(doc,
        headers=["Color", "Range", "Meaning"],
        rows=[
            ["Green", "80% and above", "Strong material match"],
            ["Amber", "50% \u2013 79%", "Moderate match"],
            ["Red", "Below 50%", "Weak match"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    add_para(doc, (
        "Hover over a row and click the eye icon next to the Match % to view a detailed Material "
        "Breakdown dialog showing exactly which materials matched."
    ))

    add_heading(doc, "Available Actions by Status", level=3)
    add_styled_table(doc,
        headers=["Status", "Actions"],
        rows=[
            ["Open", "Submit button, Excel download button"],
            ["Submitted", "Excel download button"],
            ["Draft", "Mark Submitted button, Excel download button"],
            ["Declined", "Excel download button"],
        ]
    )

    if "image7.png" in image_paths:
        add_figure(doc, image_paths["image7.png"], "RFP Table with Actions")

    # 3.4 Sync Portal
    add_heading(doc, "3.4 Sync Portal", level=2)
    add_para(doc, (
        "Click the Sync Portal button (top-right of the RFP Management section) to refresh data "
        "from the company portals. This fetches the latest RFP statuses and updates the dashboard."
    ))

    doc.add_page_break()


def build_section_4(doc, image_paths):
    """Section 4: Key Workflows"""
    add_heading(doc, "4. Key Workflows", level=1)

    # 4.1 Submit an RFP
    add_heading(doc, "4.1 Submit an RFP", level=2)
    add_para(doc, "You can submit an RFP in two ways:")

    add_para(doc, "Method 1: From the Dashboard Table", bold=True, space_after=Pt(4))
    add_numbered(doc, "Navigate to Dashboard.")
    add_numbered(doc, "Select the company tab, then the Open status tab.")
    add_numbered(doc, "Find the RFP you want to submit.")
    add_numbered(doc, "Click the Submit button in the Actions column.")
    add_numbered(doc, "The Submit RFP dialog opens with the RFP ID pre-filled.")

    add_para(doc, "Method 2: From the Sidebar Quick Action", bold=True, space_after=Pt(4))
    add_numbered(doc, "Click Submit RFP (green button) in the sidebar\u2019s Quick Actions section.")
    add_numbered(doc, "The Submit RFP dialog opens with a blank RFP ID field.")

    add_para(doc, "Submit RFP Dialog Steps:", bold=True, space_after=Pt(4))
    add_numbered(doc, " \u2013 Enter or confirm the RFP ID. The system validates it automatically: "
                      "a green checkmark appears if the RFP is found, or a red error if it is not found "
                      "(it must be downloaded first).", bold_prefix="RFP ID")
    add_numbered(doc, " \u2013 Auto-populated based on the RFP record. This field is locked when the RFP is validated.", bold_prefix="Company")
    add_numbered(doc, " (required) \u2013 Click the dashed upload area to select your filled RFP Excel file "
                      "(.xls or .xlsx).", bold_prefix="Upload Excel File")
    add_numbered(doc, " (optional) \u2013 Click to upload one or more PDF files (Technical Data Sheets). "
                      "You can remove uploaded PDFs by clicking the X button next to each file.", bold_prefix="Technical PDF Files")
    add_numbered(doc, "Click Submit RFP to start the automation.")

    add_para(doc, "The system will:", bold=True, space_after=Pt(4))
    add_bullet(doc, "Upload files to SharePoint")
    add_bullet(doc, "Navigate to the company portal")
    add_bullet(doc, "Fill in the submission form")
    add_bullet(doc, "Upload the Excel and technical files")
    add_bullet(doc, "Save the RFP as a draft on the portal")

    add_para(doc, (
        "A success notification will appear when the process starts. You can track progress via "
        "the Automation Status indicator in the sidebar footer."
    ))

    if "image8.png" in image_paths:
        add_figure(doc, image_paths["image8.png"], "Submit RFP Dialog", width=Inches(3.5))

    # 4.2 Decline an RFP
    add_heading(doc, "4.2 Decline an RFP", level=2)
    add_numbered(doc, "Click Decline RFP (red button) in the sidebar\u2019s Quick Actions section.")
    add_numbered(doc, "The Decline RFP dialog opens.")
    add_numbered(doc, " \u2013 Enter the exact RFP title. The system validates it: "
                      "a green checkmark means the RFP is found, a red error means it is not in the database.", bold_prefix="RFP Title")
    add_numbered(doc, " \u2013 Auto-populated and locked after validation.", bold_prefix="Company")
    add_numbered(doc, "Click Decline RFP (red button) to confirm.")

    add_para(doc, (
        "The system will navigate to the portal and decline the RFP. A success notification "
        "appears when the process starts."
    ))

    # 4.3 Download RFPs
    add_heading(doc, "4.3 Download RFPs", level=2)
    add_numbered(doc, "Click Download RFPs (blue button) in the sidebar\u2019s Quick Actions section.")
    add_numbered(doc, "The Download dialog opens with two options:")
    add_bullet(doc, " \u2013 Download from all configured company portals.", bold_prefix="All Companies")
    add_bullet(doc, " \u2013 Select a single company from the dropdown.", bold_prefix="Specific Company")
    add_numbered(doc, "Review the information about what the automation will do.")
    add_numbered(doc, "Click Yes, Start Download to begin.")

    add_para(doc, "The automation will:", bold=True, space_after=Pt(4))
    add_bullet(doc, "Navigate to the selected company portal(s)")
    add_bullet(doc, "Scrape RFP listings and download RFP files")
    add_bullet(doc, "Save information to the database")
    add_bullet(doc, "Send email notifications for new RFPs found")

    if "image9.png" in image_paths:
        add_figure(doc, image_paths["image9.png"], "Download RFPs Dialog", width=Inches(3.0))

    # 4.4 Download RFP Excel File
    add_heading(doc, "4.4 Download RFP Excel File", level=2)
    add_para(doc, "To download the Excel file for any individual RFP:")
    add_numbered(doc, "Find the RFP in the Dashboard table (any status tab).")
    add_numbered(doc, "Click the green Excel button in the Actions column.")
    add_numbered(doc, "The file will download to your browser\u2019s default download location.")
    add_numbered(doc, "A success notification confirms the download.")

    # 4.5 Respond to RFP Email (Adaptive Card)
    add_heading(doc, "4.5 Respond to RFP Email (Adaptive Card)", level=2)
    add_para(doc, (
        "When a new RFP is found, team members receive an email notification in Microsoft Outlook "
        "with an interactive Adaptive Card."
    ))

    add_para(doc, "How to respond:", bold=True, space_after=Pt(4))
    add_numbered(doc, "Open the email notification in Outlook. You will see:")
    add_bullet(doc, "RFP details (ID, company, deadline)")
    add_bullet(doc, "A table listing all assigned team members (your name is highlighted)")
    add_bullet(doc, "Input fields for Results and Remarks")
    add_numbered(doc, "Fill in your Results (your assessment or bid information).")
    add_numbered(doc, "Fill in your Remarks (any additional notes or comments).")
    add_numbered(doc, "Click the Submit button within the email card.")
    add_numbered(doc, "The card updates in-place to show your submission was recorded.")

    add_para(doc, "What happens next:", bold=True, space_after=Pt(4))
    add_para(doc, (
        "Once all team members have submitted their responses, a consolidated summary email is "
        "automatically sent to the configured stakeholders."
    ))

    doc.add_page_break()


def build_section_5(doc, image_paths):
    """Section 5: RFP Insights"""
    add_heading(doc, "5. RFP Insights", level=1)
    add_para(doc, (
        "Navigate to RFP Insights from the sidebar menu. This page provides a detailed, "
        "filterable view of all RFPs across all companies."
    ))

    # 5.1 Stats Overview
    add_heading(doc, "5.1 Stats Overview", level=2)
    add_para(doc, "Five stat cards at the top show:")
    add_styled_table(doc,
        headers=["Stat", "Description"],
        rows=[
            ["Total RFPs", "All RFPs in the system"],
            ["Submitted", "RFPs that have been submitted"],
            ["Declined", "RFPs that were declined"],
            ["Not Participant", "RFPs with no participation action taken"],
            ["Open", "Currently open RFPs"],
        ]
    )

    # 5.2 Advanced Filters
    add_heading(doc, "5.2 Advanced Filters", level=2)
    add_para(doc, "The filter panel lets you narrow down RFPs using multiple criteria:")
    add_styled_table(doc,
        headers=["Filter", "Options"],
        rows=[
            ["Status", "All, Open, Submitted, Declined, Not Participant"],
            ["Company", "Dropdown populated from your data"],
            ["Start Date / End Date", "Date range pickers"],
            ["Material Match", "All, Matched, Not Matched"],
            ["Keyword Match", "All, Matched, Not Matched"],
            ["Participation", "All, Participated, Not Participated, Declined"],
            ["Search", "Free-text search by RFP ID, company, or owner"],
        ]
    )

    add_para(doc, "Click Apply to filter the results. Click Reset to clear all filters.")

    # 5.3 Column Visibility
    add_heading(doc, "5.3 Column Visibility", level=2)
    add_para(doc, (
        "Click the Column Visibility dropdown (gear icon) to toggle which columns appear in the table. "
        "Your selection is saved automatically and persists across sessions."
    ))

    # 5.4 Results Table
    add_heading(doc, "5.4 Results Table", level=2)
    add_para(doc, (
        "The table displays filtered results with infinite scroll \u2014 as you scroll down, more rows load "
        "automatically. Each row shows:"
    ))
    add_bullet(doc, "Status badge with icon")
    add_bullet(doc, "Participation status")
    add_bullet(doc, "Portal link button (opens RFP on the external portal)")
    add_bullet(doc, "Download Excel button")

    if "image10.png" in image_paths:
        add_figure(doc, image_paths["image10.png"], "RFP Insights Page")

    doc.add_page_break()


def build_section_6(doc, image_paths):
    """Section 6: Material Insights"""
    add_heading(doc, "6. Material Insights", level=1)
    add_para(doc, (
        "Navigate to Material Insights from the sidebar menu. This page analyzes how your "
        "company\u2019s materials and keywords match against RFPs."
    ))

    # 6.1 Stats Overview
    add_heading(doc, "6.1 Stats Overview", level=2)
    add_para(doc, "Four stat cards show:")
    add_bullet(doc, " \u2013 Number of distinct material codes in the system", bold_prefix="Unique Materials")
    add_bullet(doc, " \u2013 Number of distinct keywords tracked", bold_prefix="Unique Keywords")
    add_bullet(doc, " \u2013 How many RFPs had at least one material or keyword match", bold_prefix="RFPs with Matches")
    add_bullet(doc, " \u2013 How many matched RFPs were submitted", bold_prefix="Submitted RFPs")

    # 6.2 Tabs
    add_heading(doc, "6.2 Tabs: Materials vs Keywords", level=2)
    add_para(doc, "Toggle between two views using the tabs:")
    add_bullet(doc, " \u2013 Analysis by material code", bold_prefix="Materials")
    add_bullet(doc, " \u2013 Analysis by keyword", bold_prefix="Keywords")

    # 6.3 Charts
    add_heading(doc, "6.3 Charts", level=2)
    add_para(doc, "Depending on the selected tab:")
    add_bullet(doc, " \u2013 Top 10 Materials (or Keywords) ranked by RFP count", bold_prefix="Bar Chart")
    add_bullet(doc, " \u2013 Keyword frequency distribution (visible on Keywords tab)", bold_prefix="Pie Chart")

    # 6.4 Filters
    add_heading(doc, "6.4 Filters", level=2)
    add_styled_table(doc,
        headers=["Filter", "Options"],
        rows=[
            ["Company", "Filter by specific company"],
            ["Participation", "All, Submitted, Declined, Open"],
            ["Search", "Search by material code or keyword"],
        ]
    )

    # 6.5 Expandable Table
    add_heading(doc, "6.5 Expandable Table", level=2)
    add_para(doc, "The main table shows materials or keywords as parent rows. Each row shows:")
    add_bullet(doc, "Code/Keyword, Description, RFP Count, Companies, Submitted Count")

    add_para(doc, "Click the expand arrow on any row to reveal the individual RFPs that matched:")
    add_bullet(doc, "RFP ID, Company, Deadline, Match Method (Exact or Keyword), Participation status")

    add_para(doc, "The table supports infinite scroll for large datasets.")

    if "image11.png" in image_paths:
        add_figure(doc, image_paths["image11.png"], "Material Insights Page")

    doc.add_page_break()


def build_section_7(doc, image_paths):
    """Section 7: Activity Logs"""
    add_heading(doc, "7. Activity Logs", level=1)
    add_para(doc, "Navigate to Activity Logs from the sidebar menu. This page shows the history of all automation runs.")

    # 7.1 Statistics Cards
    add_heading(doc, "7.1 Statistics Cards", level=2)
    add_para(doc, "Four clickable cards at the top:")
    add_bullet(doc, " \u2013 All automation runs", bold_prefix="Total Runs")
    add_bullet(doc, " \u2013 Successfully completed runs (click to filter)", bold_prefix="Completed")
    add_bullet(doc, " \u2013 Runs that encountered errors (click to filter)", bold_prefix="Failed")
    add_bullet(doc, " \u2013 Currently active runs (click to filter)", bold_prefix="Running")

    # 7.2 Controls
    add_heading(doc, "7.2 Controls", level=2)
    add_bullet(doc, " \u2013 Search by RFP ID, action, or run ID", bold_prefix="Search box")
    add_bullet(doc, " \u2013 Choose how many logs to show (50, 100, 200, or 500)", bold_prefix="Page size selector")

    # 7.3 Run Cards
    add_heading(doc, "7.3 Run Cards", level=2)
    add_para(doc, "Each automation run is displayed as a card with:")
    add_bullet(doc, "Left color bar: Green = Completed, Red = Failed, Blue = Running")
    add_bullet(doc, "Status icon: Checkmark, X, or spinner")
    add_bullet(doc, "RFP ID and status badge")
    add_bullet(doc, "Action name (e.g., Download, Submit, Decline)")
    add_bullet(doc, "Start/End time")
    add_bullet(doc, "Step progress (e.g., 5 success / 0 failed / 5 total)")
    add_bullet(doc, "View Details button")

    if "image12.png" in image_paths:
        add_figure(doc, image_paths["image12.png"], "Activity Logs Page")

    # 7.4 Run Detail Modal
    add_heading(doc, "7.4 Run Detail Modal", level=2)
    add_para(doc, "Click View Details to open a modal with three tabs:")

    add_para(doc, "Timeline Tab", bold=True, space_after=Pt(4))
    add_para(doc, "A vertical timeline showing each step of the automation:")
    add_bullet(doc, "Timestamp, action type, status, and details for each step")
    add_bullet(doc, "Color-coded dots: Green = success, Red = failure, Gray = skipped")

    add_para(doc, "Error Report Tab", bold=True, space_after=Pt(4))
    add_para(doc, "If the run encountered errors, this tab shows:")
    add_bullet(doc, "Error type and summary")
    add_bullet(doc, "Context information")
    add_bullet(doc, "Full traceback")
    add_bullet(doc, "Suggested actions to resolve")

    add_para(doc, "Screenshots Tab", bold=True, space_after=Pt(4))
    add_para(doc, (
        "If the automation captured browser screenshots during failure, they are displayed here "
        "for visual debugging."
    ))

    doc.add_page_break()


def build_section_8(doc, image_paths):
    """Section 8: Analytics"""
    add_heading(doc, "8. Analytics", level=1)
    add_para(doc, "Navigate to Analytics from the sidebar menu. This page provides visual charts and interactive drill-downs.")

    # 8.1 Key Metrics
    add_heading(doc, "8.1 Key Metrics", level=2)
    add_para(doc, "Four cards at the top:")
    add_bullet(doc, " \u2013 All RFPs in the system", bold_prefix="Total RFPs")
    add_bullet(doc, " \u2013 Submitted count with participation rate percentage", bold_prefix="Submitted")
    add_bullet(doc, " \u2013 Count and percentage of RFPs with material matches", bold_prefix="Material Matched")
    add_bullet(doc, " \u2013 Count and percentage of RFPs with keyword matches", bold_prefix="Keyword Matched")

    add_tip(doc, "Click any metric card to navigate to RFP Insights with the relevant filter applied.")

    # 8.2 Interactive Charts
    add_heading(doc, "8.2 Interactive Charts", level=2)
    add_para(doc, "The page displays four charts in a 2-column layout:")

    add_styled_table(doc,
        headers=["Chart", "Type", "What It Shows"],
        rows=[
            ["RFP Status Distribution", "Donut chart", "Breakdown by Submitted (green), Open (amber), Declined (red)"],
            ["Top 5 Companies", "Horizontal bar chart", "Companies ranked by RFP count"],
            ["Material Matching", "Donut chart", "Material Matched vs Not Matched"],
            ["Keyword Matching", "Donut chart", "Keyword Matched vs Not Matched"],
        ]
    )

    add_para(doc, (
        "Drill-Down: Click on any chart segment or bar to navigate to the RFP Insights page "
        "with the corresponding filter pre-applied."
    ))

    # 8.3 Participation by Company
    add_heading(doc, "8.3 Participation by Company", level=2)
    add_para(doc, "Below the charts, a horizontal stacked bar chart shows participation breakdown per company:")
    add_bullet(doc, "Participated (green), Not Participated (gray), Declined (red)")

    add_para(doc, "Click any segment to drill down to the corresponding filtered view.")

    if "image13.png" in image_paths:
        add_figure(doc, image_paths["image13.png"], "Analytics Dashboard")

    doc.add_page_break()


def build_section_9(doc):
    """Section 9: Profile Settings"""
    add_heading(doc, "9. Profile Settings", level=1)
    add_para(doc, "Navigate to Profile by clicking your profile icon in the header or navigating to the profile page.")

    # 9.1 Update Profile Information
    add_heading(doc, "9.1 Update Profile Information", level=2)
    add_para(doc, "The Profile Information card shows:")
    add_bullet(doc, " \u2013 Editable. Change your name and click Save Changes.", bold_prefix="Display Name")
    add_bullet(doc, " \u2013 Read-only. Contact an administrator to change your email.", bold_prefix="Email")
    add_bullet(doc, " \u2013 Optional. Add or update your phone number.", bold_prefix="Mobile Number")
    add_bullet(doc, " \u2013 Read-only. Shows your assigned role (e.g., RFP Bidder).", bold_prefix="Role")

    # 9.2 Change Password
    add_heading(doc, "9.2 Change Password", level=2)
    add_para(doc, "The Change Password card has three fields:")
    add_bullet(doc, " \u2013 Enter your existing password.", bold_prefix="Current Password")
    add_bullet(doc, " \u2013 Enter a new password that meets these requirements:", bold_prefix="New Password")
    add_bullet(doc, "Minimum 8 characters")
    add_bullet(doc, "At least one uppercase letter")
    add_bullet(doc, "At least one number")
    add_bullet(doc, " \u2013 Re-enter the new password.", bold_prefix="Confirm New Password")

    add_para(doc, "Click Change Password to update. A success notification confirms the change.")

    doc.add_page_break()


def build_section_10(doc):
    """Section 10: Automation Status Indicator"""
    add_heading(doc, "10. Automation Status Indicator", level=1)
    add_para(doc, "The bottom of the sidebar shows the Automation Status.")

    add_para(doc, "When Idle (Ready):", bold=True, space_after=Pt(4))
    add_bullet(doc, 'A green dot with the text "Ready" appears.')
    add_bullet(doc, "This means no automation is currently running.")

    add_para(doc, "When Running:", bold=True, space_after=Pt(4))
    add_bullet(doc, 'The indicator changes to yellow/amber with "Running" text and an animated pulse.')
    add_bullet(doc, "A progress bar appears showing the percentage complete.")
    add_bullet(doc, "Detailed progress shows which operation is running:")
    add_bullet(doc, "Download: Shows current/total count and the current item being processed")
    add_bullet(doc, "Submit: Shows processing status message")
    add_bullet(doc, "Decline: Shows processing status message")

    add_important(doc, (
        "While automation is running, the corresponding Quick Action button in the sidebar is "
        "disabled (grayed out) to prevent duplicate runs. You can continue using other parts of "
        "the portal normally."
    ))

    add_para(doc, (
        "When the sidebar is collapsed, the automation status is shown as a small colored dot "
        "(green = Ready, amber = Running). Hover over it to see details."
    ))

    doc.add_page_break()


def build_section_11(doc):
    """Section 11: Quick Reference"""
    add_heading(doc, "11. Quick Reference", level=1)

    # 11.1 Status Colors
    add_heading(doc, "11.1 Status Colors", level=2)
    add_styled_table(doc,
        headers=["Status", "Color", "Badge Style"],
        rows=[
            ["Open", "Amber/Orange", "Amber background, amber text"],
            ["Submitted", "Green", "Green background, green text"],
            ["Draft / Saved Draft", "Gray", "Gray background, gray text"],
            ["Declined", "Red", "Red/Rose background, red text"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 11.2 Match % Thresholds
    add_heading(doc, "11.2 Match % Thresholds", level=2)
    add_styled_table(doc,
        headers=["Range", "Color", "Meaning"],
        rows=[
            ["80% and above", "Green", "Strong match \u2013 high relevance"],
            ["50% \u2013 79%", "Amber", "Moderate match \u2013 review recommended"],
            ["Below 50%", "Red", "Weak match \u2013 low relevance"],
            ["No data", "Gray dash", "No material match data available"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 11.3 Quick Action Buttons
    add_heading(doc, "11.3 Quick Action Buttons", level=2)
    add_styled_table(doc,
        headers=["Button", "Color", "Action"],
        rows=[
            ["Download RFPs", "Blue gradient", "Opens download dialog to scrape RFPs from portals"],
            ["Submit RFP", "Green gradient", "Opens submit dialog to upload files and submit"],
            ["Decline RFP", "Red outline", "Opens decline dialog to decline participation"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 11.4 Automation Run Statuses
    add_heading(doc, "11.4 Automation Run Statuses", level=2)
    add_styled_table(doc,
        headers=["Status", "Indicator", "Meaning"],
        rows=[
            ["Completed", "Green bar + checkmark", "Run finished successfully"],
            ["Failed", "Red bar + X icon", "Run encountered a fatal error"],
            ["Running", "Blue bar + spinner", "Run is currently in progress"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 11.5 Supported Companies
    add_heading(doc, "11.5 Supported Companies", level=2)
    add_styled_table(doc,
        headers=["Company", "Portal"],
        rows=[
            ["Saudi Energy", "SEC procurement portal"],
            ["Aramco e-Marketplace", "Aramco procurement portal"],
            ["SABIC \u2013 Saudi Basic Industries Corp.", "SABIC procurement portal"],
            ["HADEED \u2013 RAJHI STEEL", "HADEED procurement portal"],
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 11.6 Tips & Shortcuts (NEW)
    add_heading(doc, "11.6 Tips and Shortcuts", level=2)
    add_styled_table(doc,
        headers=["Tip / Shortcut", "Description"],
        rows=[
            ["Click metric cards", "Navigate directly to filtered RFP Insights"],
            ["Click chart segments", "Drill down to the corresponding filtered view"],
            ["Hover on Match % eye icon", "View detailed material breakdown for an RFP"],
            ["Column Visibility gear icon", "Toggle which columns appear in the RFP Insights table"],
            ["Sidebar collapse arrow", "Toggle between full and compact sidebar view"],
            ["Infinite scroll", "Scroll down to automatically load more rows in tables"],
        ]
    )

    doc.add_page_break()


def build_section_12(doc):
    """Section 12: Troubleshooting & FAQ"""
    add_heading(doc, "12. Troubleshooting and FAQ", level=1)

    # FAQ 1
    add_heading(doc, "I can\u2019t log in", level=2)
    add_bullet(doc, " Make sure you\u2019re using the correct email and password.", bold_prefix="Check your credentials:")
    add_bullet(doc, " After 5 failed login attempts within 5 minutes, your account is locked "
                    "for 30 minutes. Wait and try again, or contact your administrator to unlock it.", bold_prefix="Account locked:")
    add_bullet(doc, " Passwords must be changed every 90 days. Use the Forgot Password link to reset.", bold_prefix="Password expired:")

    # FAQ 2
    add_heading(doc, "The dashboard shows no data", level=2)
    add_bullet(doc, " Click the Download RFPs button in the sidebar to download RFPs from the portals.", bold_prefix="First time?")
    add_bullet(doc, " Click the Sync Portal button in the RFP Management section header to refresh.", bold_prefix="Data stale?")
    add_bullet(doc, "Check Activity Logs to see if a download automation has run recently.")

    # FAQ 3
    add_heading(doc, "RFP submission failed", level=2)
    add_bullet(doc, "Go to Activity Logs and find the failed run.")
    add_bullet(doc, "Click View Details and check the Error Report tab for the error message.")
    add_bullet(doc, "Common causes: portal was unavailable, file format was incorrect, or session timed out.")
    add_bullet(doc, "If the issue persists, contact your administrator.")

    # FAQ 4
    add_heading(doc, "I don\u2019t see certain menu items", level=2)
    add_bullet(doc, (
        "Menu items are based on your role and permissions. RFP Bidder users may not see admin "
        "sections like Users, Roles, or Audit Logs."
    ))
    add_bullet(doc, "If you need access, contact your administrator to update your role permissions.")

    # FAQ 5
    add_heading(doc, "My session expired", level=2)
    add_bullet(doc, "Sessions time out after 2 hours of activity or 30 minutes of inactivity.")
    add_bullet(doc, "You will be redirected to the login page. Simply log in again to continue.")
    add_bullet(doc, (
        "Your unsaved form data (e.g., in a Submit dialog) will be lost \u2014 complete submissions promptly."
    ))

    # FAQ 6
    add_heading(doc, "The Submit/Decline button is grayed out", level=2)
    add_bullet(doc, (
        "This means the corresponding automation is already running. Check the Automation Status "
        "indicator at the bottom of the sidebar."
    ))
    add_bullet(doc, "Wait for the current operation to complete before starting a new one.")

    # FAQ 7
    add_heading(doc, "I entered an RFP ID but got an error in the Submit/Decline dialog", level=2)
    add_bullet(doc, (
        'The message "RFP not found in database" means the RFP must be downloaded first before it '
        "can be submitted or declined."
    ))
    add_bullet(doc, "Use the Download RFPs action to download it, then try again.")

    # FAQ 8
    add_heading(doc, "Excel download is not working", level=2)
    add_bullet(doc, "Ensure your browser allows file downloads and pop-ups from the portal URL.")
    add_bullet(doc, (
        "Check if the RFP has an associated Excel file \u2014 some newly downloaded RFPs may not have files yet."
    ))

    # FAQ 9
    add_heading(doc, "How do I know if my Adaptive Card response was saved?", level=2)
    add_bullet(doc, (
        "After clicking Submit in the Outlook Adaptive Card, the card should update to show a confirmation message."
    ))
    add_bullet(doc, "If the card doesn\u2019t update, check your internet connection and try again.")
    add_bullet(doc, "Your responses are saved in the system \u2014 check with your administrator if unsure.")

    # FAQ 10 (NEW)
    add_heading(doc, "How do I change my display name?", level=2)
    add_bullet(doc, (
        "Navigate to Profile Settings (click your profile icon in the top-right header). "
        "Edit the Display Name field and click Save Changes."
    ))

    # FAQ 11 (NEW)
    add_heading(doc, "Can I use the portal on a mobile device?", level=2)
    add_bullet(doc, (
        "The portal is optimized for desktop browsers. While it may work on mobile browsers, "
        "the full experience is best on Google Chrome (latest version) on a desktop or laptop."
    ))

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_para(doc, "For additional support, contact your system administrator.", bold=True)

    doc.add_page_break()


def build_glossary(doc):
    """Appendix A: Glossary of Terms"""
    add_heading(doc, "Appendix A: Glossary", level=1)

    add_styled_table(doc,
        headers=["Term", "Definition"],
        rows=[
            ["RFP", "Request for Proposal \u2013 a formal document issued by a company inviting suppliers to submit bids for products or services."],
            ["Adaptive Card", "An interactive email component in Microsoft Outlook that allows users to respond to RFP notifications directly within the email."],
            ["Match %", "The percentage of materials in an RFP that match your company\u2019s product catalog, indicating relevance."],
            ["Automation Run", "A single execution of the portal automation (download, submit, or decline) tracked in Activity Logs."],
            ["SharePoint", "Microsoft cloud storage where RFP files (Excel, PDFs) are uploaded during submission."],
            ["Portal", "An external procurement website of a company (e.g., SEC, Aramco) where RFPs are published and managed."],
            ["TDS", "Technical Data Sheet \u2013 PDF documentation describing product specifications, uploaded during RFP submission."],
            ["Sync", "The process of refreshing data from external company portals to update the dashboard with the latest RFP statuses."],
            ["Draft", "An RFP submission saved on the company portal but not yet finalized or officially submitted."],
            ["Quick Action", "Sidebar buttons (Download, Submit, Decline) that provide one-click access to common workflows."],
        ],
        col_widths=[Inches(1.8), Inches(4.7)]
    )

    doc.add_page_break()


def build_back_cover(doc, logo_path):
    """Build a simple back cover page."""
    # Spacers
    for _ in range(8):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # Logo
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_path):
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Inches(2.5))

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Company name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(DOC_SUBTITLE)
    run_name.font.name = FONT_HEADING
    run_name.font.size = Pt(16)
    run_name.font.color.rgb = BAHRA_DARK_GRAY
    run_name.bold = True

    # Copyright
    p_copy = doc.add_paragraph()
    p_copy.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_copy = p_copy.add_run("\u00A9 2026 Bahra Electric. All rights reserved.")
    run_copy.font.name = FONT_BODY
    run_copy.font.size = Pt(10)
    run_copy.font.color.rgb = BAHRA_LIGHT_GRAY


# =============================================================================
# MAIN ORCHESTRATOR
# =============================================================================

def generate_manual():
    """Main entry point to generate the professional user manual."""
    print("=" * 60)
    print("  Bahra Electric - User Manual Generator")
    print("=" * 60)

    # Step 1: Extract images from original document
    print("\n[1/4] Extracting images from original document...")
    if not os.path.exists(ORIGINAL_DOCX):
        print(f"  ERROR: Original document not found at {ORIGINAL_DOCX}")
        sys.exit(1)

    images = extract_images_from_docx(ORIGINAL_DOCX)
    print(f"  Found {len(images)} images: {', '.join(sorted(images.keys()))}")

    # Step 2: Save images to temp directory
    temp_dir = tempfile.mkdtemp(prefix="manual_images_")
    image_paths = {}
    for filename, data in images.items():
        path = os.path.join(temp_dir, filename)
        with open(path, 'wb') as f:
            f.write(data)
        image_paths[filename] = path
    print(f"  Saved images to temp directory: {temp_dir}")

    # Step 3: Create new document
    print("\n[2/4] Building document structure...")
    doc = Document()

    # Configure page layout
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.75)

    # Configure styles
    configure_styles(doc)

    # Step 4: Build all sections
    print("  Building cover page...")
    build_cover_page(doc, LOGO_PATH)

    print("  Building front matter...")
    build_confidentiality_notice(doc)
    build_document_control(doc)
    build_table_of_contents(doc)

    print("  Building Section 1: Introduction...")
    build_section_1(doc)

    print("  Building Section 2: Getting Started...")
    build_section_2(doc, image_paths)

    print("  Building Section 3: Dashboard...")
    build_section_3(doc, image_paths)

    print("  Building Section 4: Key Workflows...")
    build_section_4(doc, image_paths)

    print("  Building Section 5: RFP Insights...")
    build_section_5(doc, image_paths)

    print("  Building Section 6: Material Insights...")
    build_section_6(doc, image_paths)

    print("  Building Section 7: Activity Logs...")
    build_section_7(doc, image_paths)

    print("  Building Section 8: Analytics...")
    build_section_8(doc, image_paths)

    print("  Building Section 9: Profile Settings...")
    build_section_9(doc)

    print("  Building Section 10: Automation Status...")
    build_section_10(doc)

    print("  Building Section 11: Quick Reference...")
    build_section_11(doc)

    print("  Building Section 12: Troubleshooting & FAQ...")
    build_section_12(doc)

    print("  Building Glossary...")
    build_glossary(doc)

    print("  Building Back Cover...")
    build_back_cover(doc, LOGO_PATH)

    # Step 5: Setup headers and footers
    print("\n[3/4] Configuring headers and footers...")
    setup_headers_footers(doc)

    # Step 6: Save
    print("\n[4/4] Saving document...")
    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"  Saved to: {OUTPUT_DOCX}")

    # Cleanup temp images
    shutil.rmtree(temp_dir, ignore_errors=True)

    print("\n" + "=" * 60)
    print("  DONE! Open the document in Microsoft Word.")
    print("  Right-click the Table of Contents > 'Update Field'")
    print("  to populate page numbers.")
    print("=" * 60)


if __name__ == "__main__":
    generate_manual()
